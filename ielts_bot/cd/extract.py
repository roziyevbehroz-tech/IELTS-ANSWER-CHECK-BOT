"""Fayl (yoki matn)dan toza matn ajratish — AI'siz.

Qo'llab-quvvatlanadi:
    .txt / oddiy matn
    .docx           — python-docx (bo'lmasa, zip+XML fallback)
    .pdf            — pdfplumber yoki PyPDF2 (bo'lmasa, xato beradi)
    .doc (eski)     — best-effort (antiword yoki matn ajratish); ko'pincha
                       foydalanuvchidan .docx/.pdf so'raladi.

Barcha funksiyalar oddiy `str` qaytaradi. Hech qanday tarmoq/AI ishlatilmaydi.
"""

from __future__ import annotations

import io
import re
import zipfile
from typing import Optional


class ExtractError(Exception):
    """Fayldan matn ajratib bo'lmaganda."""


def extract_text(data: bytes, filename: str = "") -> str:
    """Fayl baytlaridan matn ajratadi. Kengaytmaga qarab usul tanlaydi."""
    name = (filename or "").lower()
    if name.endswith(".pdf") or data[:5] == b"%PDF-":
        return _clean(_from_pdf(data))
    if name.endswith(".docx") or data[:2] == b"PK":
        return _clean(_from_docx(data))
    if name.endswith(".doc"):
        return _clean(_from_doc(data))
    # oddiy matn
    return _clean(_from_txt(data))


# ------------------------------- TXT -------------------------------

def _from_txt(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "cp1251", "latin-1"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="ignore")


# ------------------------------- DOCX ------------------------------

def _from_docx(data: bytes) -> str:
    try:
        import docx  # type: ignore

        doc = docx.Document(io.BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            parts.append(para.text)
        # jadval matnini ham qo'shamiz (table completion manbasi bo'lishi mumkin)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except ImportError:
        return _docx_via_zip(data)
    except Exception as e:  # noqa: BLE001
        # python-docx yiqilsa, xom XML fallback
        try:
            return _docx_via_zip(data)
        except Exception:
            raise ExtractError(f"docx o'qib bo'lmadi: {e}") from e


def _docx_via_zip(data: bytes) -> str:
    """python-docx bo'lmasa: document.xml dan matn tortib olish."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
    except Exception as e:  # noqa: BLE001
        raise ExtractError(f"docx tuzilmasi buzuq: {e}") from e
    # <w:p> -> qator, <w:t> -> matn
    xml = re.sub(r"</w:p>", "\n", xml)
    xml = re.sub(r"<w:tab/>", "\t", xml)
    xml = re.sub(r"<[^>]+>", "", xml)
    # XML entity'larni ochamiz
    for a, b in (("&amp;", "&"), ("&lt;", "<"), ("&gt;", ">"),
                 ("&quot;", '"'), ("&apos;", "'")):
        xml = xml.replace(a, b)
    return xml


# ------------------------------- PDF -------------------------------

def _from_pdf(data: bytes) -> str:
    # 1) pdfplumber (eng aniq layout)
    try:
        import pdfplumber  # type: ignore

        parts = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                parts.append(txt)
        text = "\n".join(parts)
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception:  # noqa: BLE001
        pass

    # 2) PyPDF2 / pypdf
    for modname in ("pypdf", "PyPDF2"):
        try:
            mod = __import__(modname)
            reader = mod.PdfReader(io.BytesIO(data))
            parts = [(page.extract_text() or "") for page in reader.pages]
            text = "\n".join(parts)
            if text.strip():
                return text
        except ImportError:
            continue
        except Exception:  # noqa: BLE001
            continue

    raise ExtractError(
        "PDF o'qish uchun kutubxona topilmadi yoki matn ajratib bo'lmadi "
        "(skanerlangan PDF bo'lishi mumkin). Iltimos matnni .docx yoki oddiy "
        "matn ko'rinishida yuboring."
    )


# ------------------------------- DOC (eski) ------------------------

def _from_doc(data: bytes) -> str:
    # Eski binary .doc — ishonchli parse murakkab. antiword bo'lsa ishlatamiz.
    import shutil
    import subprocess
    import tempfile

    if shutil.which("antiword"):
        try:
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=True) as f:
                f.write(data)
                f.flush()
                out = subprocess.run(
                    ["antiword", f.name], capture_output=True, timeout=30
                )
                if out.returncode == 0 and out.stdout.strip():
                    return out.stdout.decode("utf-8", errors="ignore")
        except Exception:  # noqa: BLE001
            pass
    # Fallback: o'qiladigan ASCII bo'laklarini tortib olish (past sifat)
    text = data.decode("latin-1", errors="ignore")
    text = re.sub(r"[^\x09\x0a\x0d\x20-\x7e]+", " ", text)
    text = re.sub(r"\s{3,}", "\n", text)
    if len(text.strip()) < 40:
        raise ExtractError(
            "Eski .doc formatini o'qib bo'lmadi. Iltimos faylni .docx yoki "
            ".pdf ko'rinishida saqlab qayta yuboring."
        )
    return text


# ------------------------------- tozalash --------------------------

def _clean(text: str) -> str:
    """Umumiy tozalash: satrlarni normallashtirish, ortiqcha bo'shliqlar."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace(" ", " ")   # non-breaking space
    text = text.replace("﻿", "")    # BOM
    # 3+ bo'sh qatorni 2 taga
    text = re.sub(r"\n{3,}", "\n\n", text)
    # qator oxiridagi bo'shliqlar
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return text.strip()
